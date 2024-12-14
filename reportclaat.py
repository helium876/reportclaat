#!/usr/bin/env python3

import os
import socket
import ipaddress
import validators
import nmap
import subprocess
import dns.resolver
from typing import List, Dict, Set, Optional
from collections import defaultdict
from docx import Document
from docx.shared import Inches
from PIL import Image
import tempfile
import time
import requests
import urllib3
import sys
import asyncio
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor, as_completed
from rich.console import Console
from rich.progress import Progress, SpinnerColumn, TextColumn, BarColumn, TaskProgressColumn
from rich.panel import Panel
from rich.live import Live
from rich.table import Table

# Add Sublist3r to Python path
sys.path.append('Sublist3r')
import sublist3r

# Initialize rich console
console = Console()

class DomainAnalyzer:
    def __init__(self, input_file: str, max_workers: int = None, timeout: int = 30, quiet: bool = False):
        self.input_file = input_file
        self.timeout = timeout
        self.quiet = quiet
        self.ip_data: Dict[str, Dict] = defaultdict(lambda: {
            'domains': set(),
            'ports': set(),
            'services': set(),
            'screenshots': set()
        })
        self.temp_dir = tempfile.mkdtemp()
        self.nm = nmap.PortScanner()
        self.max_workers = max_workers or min(32, (os.cpu_count() or 1) * 4)
        self.executor = ThreadPoolExecutor(max_workers=self.max_workers)
        self.scan_semaphore = asyncio.Semaphore(10)
        self.screenshot_semaphore = asyncio.Semaphore(5)
        
        # Statistics for progress tracking
        self.stats = {
            'total_domains': 0,
            'processed_domains': 0,
            'total_subdomains': 0,
            'processed_subdomains': 0,
            'successful_screenshots': 0,
            'failed_screenshots': 0,
            'successful_scans': 0,
            'failed_scans': 0
        }
        
    def log(self, message: str, level: str = "info"):
        """Log a message unless quiet mode is enabled"""
        if not self.quiet:
            if level == "error":
                console.print(f"[red]{message}[/red]")
            elif level == "warning":
                console.print(f"[yellow]{message}[/yellow]")
            elif level == "success":
                console.print(f"[green]{message}[/green]")
            else:
                console.print(message)

    async def process_domain(self, domain: str, progress) -> Set[str]:
        task_id = progress.add_task(f"[cyan]Enumerating {domain}", total=None)
        try:
            subdomains = await asyncio.get_event_loop().run_in_executor(
                self.executor, self.get_subdomains, domain
            )
            progress.update(task_id, completed=True, description=f"[green]✓ {domain} ({len(subdomains)} subdomains)")
            return subdomains
        except Exception as e:
            progress.update(task_id, completed=True, description=f"[red]✗ {domain} (error: {str(e)})")
            return set()

    async def process_subdomain(self, subdomain: str, ip_addresses: Set[str], progress):
        task_id = progress.add_task(
            f"[cyan]Processing {subdomain}", 
            total=100,
            completed=0,
            visible=True
        )
        
        try:
            # Resolve IP (20%)
            progress.update(task_id, advance=20, description=f"[cyan]Resolving IP for {subdomain}")
            ip = await asyncio.get_event_loop().run_in_executor(
                self.executor, self.get_ip_for_domain, subdomain
            )
            
            if ip:
                # Add to IP data (10%)
                progress.update(task_id, advance=10, description=f"[cyan]Found IP {ip} for {subdomain}")
                self.ip_data[ip]['domains'].add(subdomain)
                
                # Port scanning (40%)
                if ip not in ip_addresses:
                    async with self.scan_semaphore:
                        progress.update(task_id, advance=20, description=f"[cyan]Scanning ports for {subdomain}")
                        await asyncio.get_event_loop().run_in_executor(
                            self.executor, self.scan_ports, ip
                        )
                        progress.update(task_id, advance=20)
                        ip_addresses.add(ip)
                else:
                    progress.update(task_id, advance=40, description=f"[cyan]Skipping port scan for {subdomain} (IP already scanned)")
                
                # Screenshot (30%)
                async with self.screenshot_semaphore:
                    progress.update(task_id, advance=15, description=f"[cyan]Taking screenshot of {subdomain}")
                    await asyncio.get_event_loop().run_in_executor(
                        self.executor, self.take_screenshot, subdomain, ip
                    )
                    progress.update(task_id, advance=15)
                
                progress.update(task_id, description=f"[green]✓ {subdomain}", completed=100)
            else:
                progress.update(task_id, description=f"[yellow]! {subdomain} (No IP)", completed=100)
                
        except Exception as e:
            progress.update(task_id, description=f"[red]✗ {subdomain} (error: {str(e)})", completed=100)
        finally:
            # Ensure task is marked as completed even if there's an error
            if not progress.tasks[task_id].completed:
                progress.update(task_id, completed=100)

    async def analyze_async(self):
        # Read input file
        ip_addresses, domains = self.read_input_file()
        self.stats['total_domains'] = len(domains)
        
        console.print(Panel(f"[bold]Analysis Started[/bold]\n"
                          f"Domains to process: {len(domains)}\n"
                          f"IP addresses to process: {len(ip_addresses)}",
                          title="Domain Analyzer", border_style="blue"))

        # Create progress tracking instance
        progress = Progress(
            SpinnerColumn(),
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TaskProgressColumn(),
            TextColumn("[cyan]{task.completed}/{task.total}"),
            TextColumn("[yellow]{task.elapsed:.2f}s"),
            TextColumn("[green]{task.fields[time_remaining]}" if "time_remaining" in "{task.fields}" else ""),
            console=console,
            expand=True,
            transient=False  # Keep failed tasks visible
        )

        with progress:
            # Enumerate subdomains
            console.print("\n[bold cyan]Starting subdomain enumeration...[/bold cyan]")
            subdomain_task = progress.add_task(
                "[cyan]Enumerating subdomains", 
                total=len(domains),
                completed=0,
                time_remaining="Calculating..."
            )
            
            # Process domains and collect subdomains
            all_subdomains = set()
            for domain in domains:
                try:
                    subdomains = await self.process_domain(domain, progress)
                    all_subdomains.update(subdomains)
                    progress.update(subdomain_task, advance=1)
                except Exception as e:
                    console.print(f"[red]Error processing domain {domain}: {str(e)}[/red]")
                    progress.update(subdomain_task, advance=1)
            
            progress.update(subdomain_task, completed=len(domains))
            self.stats['total_subdomains'] = len(all_subdomains)
            console.print(f"\n[bold cyan]Found {len(all_subdomains)} unique subdomains[/bold cyan]")

            # Process subdomains
            if all_subdomains:
                console.print("\n[bold cyan]Processing subdomains...[/bold cyan]")
                processing_task = progress.add_task(
                    "[cyan]Overall subdomain processing", 
                    total=len(all_subdomains),
                    completed=0,
                    time_remaining="Calculating..."
                )
                
                # Process subdomains in chunks to avoid overwhelming the system
                chunk_size = 5
                for i in range(0, len(all_subdomains), chunk_size):
                    chunk = list(all_subdomains)[i:i + chunk_size]
                    tasks = [self.process_subdomain(subdomain, ip_addresses, progress) 
                            for subdomain in chunk]
                    
                    await asyncio.gather(*tasks)
                    progress.update(processing_task, advance=len(chunk))
                    
                    # Update ETA
                    if i + chunk_size < len(all_subdomains):
                        remaining = len(all_subdomains) - (i + chunk_size)
                        time_per_chunk = progress.tasks[processing_task].elapsed / (i + chunk_size)
                        eta_seconds = time_per_chunk * remaining
                        time_remaining = f"~{int(eta_seconds/60)}m {int(eta_seconds%60)}s remaining"
                        progress.update(processing_task, time_remaining=time_remaining)
                
                progress.update(processing_task, completed=len(all_subdomains))

            # Process remaining IPs
            remaining_ips = [ip for ip in ip_addresses 
                           if ip not in [ip for ip in self.ip_data.keys()]]
            
            if remaining_ips:
                console.print(f"\n[bold yellow]Processing {len(remaining_ips)} remaining IP addresses[/bold yellow]")
                ip_task = progress.add_task(
                    "[cyan]Processing IPs", 
                    total=len(remaining_ips),
                    completed=0,
                    time_remaining="Calculating..."
                )
                
                for ip in remaining_ips:
                    try:
                        await self.process_ip(ip, progress)
                        progress.update(ip_task, advance=1)
                    except Exception as e:
                        console.print(f"[red]Error processing IP {ip}: {str(e)}[/red]")
                        progress.update(ip_task, advance=1)
                
                progress.update(ip_task, completed=len(remaining_ips))

        # Print final summary
        table = Table(title="Analysis Summary", show_header=True, header_style="bold magenta")
        table.add_column("Category", style="cyan")
        table.add_column("Count", justify="right", style="green")
        
        table.add_row("Total Domains Processed", str(self.stats['total_domains']))
        table.add_row("Total Subdomains Found", str(self.stats['total_subdomains']))
        table.add_row("Total IPs Processed", str(len(self.ip_data)))
        table.add_row("Successful Screenshots", str(sum(len(data['screenshots']) for data in self.ip_data.values())))
        table.add_row("Total Open Ports", str(sum(len(data['ports']) for data in self.ip_data.values())))
        
        console.print("\n")
        console.print(table)

        # Create report
        console.print("\n[bold green]Generating report...[/bold green]")
        self.create_word_document()

    async def process_ip(self, ip: str, progress) -> None:
        task_id = progress.add_task(f"[cyan]Processing IP {ip}", total=100)
        try:
            hostname = socket.gethostbyaddr(ip)[0]
            self.ip_data[ip]['domains'].add(hostname)
            progress.update(task_id, advance=50)
            
            async with self.scan_semaphore:
                await asyncio.get_event_loop().run_in_executor(
                    self.executor, self.scan_ports, ip
                )
            progress.update(task_id, description=f"[green]✓ {ip}", completed=True)
            
        except Exception as e:
            progress.update(task_id, description=f"[red]✗ {ip} (error: {str(e)})", completed=True)

    def read_input_file(self) -> tuple[Set[str], Set[str]]:
        ip_addresses = set()
        domains = set()
        
        console.print("[bold blue]Reading input file...[/bold blue]")
        with open(self.input_file, 'r') as f:
            for line in f:
                entry = line.strip()
                if not entry:
                    continue
                    
                if self.is_valid_ip(entry):
                    ip_addresses.add(entry)
                elif validators.domain(entry):
                    domains.add(entry)
        
        return ip_addresses, domains

    @staticmethod
    def is_valid_ip(ip: str) -> bool:
        try:
            ipaddress.ip_address(ip)
            return True
        except ValueError:
            return False

    def get_subdomains(self, domain: str) -> Set[str]:
        subdomains = set()
        try:
            # Add the main domain
            subdomains.add(domain)
            
            # Redirect stdout to capture Sublist3r output
            import io
            from contextlib import redirect_stdout
            output = io.StringIO()
            
            with redirect_stdout(output):
                try:
                    # Use Sublist3r with minimal reliable engines and timeout
                    found_subdomains = sublist3r.main(
                        domain,
                        20,  # Reduced number of threads
                        None,  # Output file
                        ports=None,
                        silent=True,
                        verbose=False,
                        enable_bruteforce=False,
                        engines=['google', 'bing', 'dnsdumpster']  # Most reliable engines only
                    )
                except Exception as e:
                    console.print(f"[yellow]Warning: Sublist3r encountered an error: {str(e)}[/yellow]")
                    found_subdomains = []
            
            # Process captured output to extract additional subdomains
            output_text = output.getvalue()
            
            # Look for subdomains in the output that might have been missed
            import re
            extra_subdomains = re.findall(r'(?:[a-zA-Z0-9](?:[a-zA-Z0-9\-]{0,61}[a-zA-Z0-9])?\.)+[a-zA-Z]{2,}', output_text)
            
            # Add found subdomains from Sublist3r
            if found_subdomains:
                for subdomain in found_subdomains:
                    if validators.domain(subdomain):
                        subdomains.add(subdomain.lower())
            
            # Add extra subdomains found in output
            for subdomain in extra_subdomains:
                if validators.domain(subdomain) and subdomain.endswith(domain):
                    subdomains.add(subdomain.lower())
            
            # Try common prefixes first (faster than DNS)
            common_prefixes = ['www', 'mail', 'remote', 'blog', 'webmail', 'server',
                             'ns1', 'ns2', 'smtp', 'secure', 'vpn', 'api', 'dev',
                             'staging', 'app', 'admin', 'portal']
            
            for prefix in common_prefixes:
                candidate = f"{prefix}.{domain}"
                try:
                    if socket.gethostbyname(candidate):
                        subdomains.add(candidate)
                except:
                    continue
            
            # Try DNS zone transfer with timeout
            try:
                import dns.zone
                import dns.query
                import dns.resolver
                
                # Set shorter timeout for DNS operations
                dns.resolver.default_resolver = dns.resolver.Resolver(configure=False)
                dns.resolver.default_resolver.nameservers = ['8.8.8.8', '8.8.4.4']  # Use Google DNS
                dns.resolver.default_resolver.timeout = 3
                dns.resolver.default_resolver.lifetime = 3
                
                try:
                    # Get NS records with timeout
                    ns_records = dns.resolver.resolve(domain, 'NS', lifetime=3)
                    for ns in ns_records:
                        try:
                            # Attempt zone transfer with short timeout
                            zone = dns.zone.from_xfr(dns.query.xfr(str(ns), domain, lifetime=3))
                            if zone:
                                for name, node in zone.nodes.items():
                                    subdomain = str(name) + '.' + domain
                                    if subdomain.startswith('@'):
                                        subdomain = domain
                                    if validators.domain(subdomain):
                                        subdomains.add(subdomain.lower())
                        except Exception:
                            continue
                except Exception:
                    pass
            except Exception as e:
                console.print(f"[yellow]Warning: AXFR attempt failed: {str(e)}[/yellow]")
            
            if not subdomains:
                console.print(f"[yellow]Warning: No subdomains found for {domain}, using main domain only[/yellow]")
                subdomains.add(domain)
            
            return subdomains
            
        except Exception as e:
            console.print(f"[red]Error enumerating subdomains for {domain}: {str(e)}[/red]")
            # Return at least the main domain if everything fails
            return {domain}

    def get_ip_for_domain(self, domain: str) -> Optional[str]:
        try:
            return socket.gethostbyname(domain)
        except socket.gaierror:
            console.print(f"[yellow]Could not resolve IP for {domain}[/yellow]")
            return None

    def scan_ports(self, ip: str):
        try:
            # Common ports to scan - web, mail, ftp, ssh, etc.
            common_ports = (
                "80,443,8080,8443"  # Focus on web ports first
            )
            
            # Service name mapping for cleaner output
            service_names = {
                'http': 'HTTP',
                'https': 'HTTPS',
                'ssh': 'SSH',
                'ftp': 'FTP',
                'smtp': 'SMTP',
                'imap': 'IMAP',
                'pop3': 'POP3',
                'mysql': 'MySQL',
                'ms-sql': 'MSSQL',
                'rdp': 'RDP',
                'vnc': 'VNC',
                'telnet': 'Telnet'
            }
            
            # Focused scan configuration
            scan_args = f'-sT -T4 --max-retries=1 -Pn --host-timeout=30s --min-rate=300 -p{common_ports}'
            
            # Redirect nmap output
            import io
            from contextlib import redirect_stdout, redirect_stderr
            
            # Capture both stdout and stderr
            with redirect_stdout(io.StringIO()), redirect_stderr(io.StringIO()):
                console.print(f"[cyan]Scanning web ports for {ip}...[/cyan]")
                self.nm.scan(ip, arguments=scan_args)
            
            if ip in self.nm.all_hosts():
                for proto in self.nm[ip].all_protocols():
                    ports = self.nm[ip][proto].keys()
                    for port in ports:
                        port_info = self.nm[ip][proto][port]
                        # Only add port if state is 'open' and we have a confirmed service
                        if port_info['state'] == 'open' and port_info.get('name') != 'unknown':
                            # Additional validation for common false positives
                            service_name = port_info.get('name', '').lower()
                            
                            # Skip likely false positives
                            if port == 8080 and not self._validate_web_service(ip, port):
                                continue
                                
                            self.ip_data[ip]['ports'].add(port)
                            
                            # Get clean service name only
                            clean_service = service_names.get(service_name, service_name.upper())
                            self.ip_data[ip]['services'].add(clean_service)
                            
            # If we found web ports, we're done. Otherwise, try other common ports
            if not any(port in self.ip_data[ip]['ports'] for port in [80, 443, 8080, 8443]):
                other_ports = "21,22,23,25,110,143,3306,3389,5900"  # Reduced to most common
                scan_args = f'-sT -T4 --max-retries=1 -Pn --host-timeout=30s --min-rate=300 -p{other_ports}'
                
                with redirect_stdout(io.StringIO()), redirect_stderr(io.StringIO()):
                    console.print(f"[cyan]Scanning other common ports for {ip}...[/cyan]")
                    self.nm.scan(ip, arguments=scan_args)
                
                if ip in self.nm.all_hosts():
                    for proto in self.nm[ip].all_protocols():
                        ports = self.nm[ip][proto].keys()
                        for port in ports:
                            port_info = self.nm[ip][proto][port]
                            if port_info['state'] == 'open' and port_info.get('name') != 'unknown':
                                service_name = port_info.get('name', '').lower()
                                
                                # Skip likely false positives
                                if (port == 53 and not self._validate_dns_service(ip)):
                                    continue
                                    
                                self.ip_data[ip]['ports'].add(port)
                                
                                # Get clean service name only
                                clean_service = service_names.get(service_name, service_name.upper())
                                self.ip_data[ip]['services'].add(clean_service)
                            
        except Exception as e:
            console.print(f"[red]Error scanning ports for {ip}: {str(e)}[/red]")
            if str(e).startswith('Nmap not found'):
                console.print("[yellow]Please ensure nmap is installed: brew install nmap[/yellow]")

    def _validate_dns_service(self, ip: str) -> bool:
        """Validate if DNS service is really running by attempting a query"""
        try:
            resolver = dns.resolver.Resolver()
            resolver.nameservers = [ip]
            resolver.timeout = 2
            resolver.lifetime = 2
            # Try to resolve a common domain
            resolver.resolve('google.com', 'A')
            return True
        except:
            return False

    def _validate_web_service(self, ip: str, port: int) -> bool:
        """Validate if web service is really running by attempting a connection"""
        try:
            url = f"http://{ip}:{port}"
            response = requests.get(url, timeout=5, verify=False)
            return response.status_code < 500  # Accept any response that's not a server error
        except:
            return False

    def take_screenshot(self, domain: str, ip: str):
        try:
            from playwright.sync_api import sync_playwright, TimeoutError as PlaywrightTimeout
            console.print(f"[yellow]Taking screenshot of {domain}...[/yellow]")
            
            with sync_playwright() as p:
                try:
                    # Launch browser with specific options
                    browser = p.chromium.launch(
                        headless=True,
                        args=[
                            '--no-sandbox',
                            '--disable-setuid-sandbox',
                            '--disable-dev-shm-usage',
                            '--disable-gpu',
                            '--window-size=1920,1080',
                        ]
                    )
                    
                    # Create context with specific viewport
                    context = browser.new_context(
                        viewport={'width': 1920, 'height': 1080},
                        ignore_https_errors=True,
                        user_agent='Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36'
                    )
                    
                    # Create new page
                    page = context.new_page()
                    page.set_default_timeout(15000)  # 15 seconds
                    
                    success = False
                    for protocol in ['https', 'http']:
                        if success:
                            break
                            
                        url = f"{protocol}://{domain}"
                        console.print(f"[yellow]Attempting {url}[/yellow]")
                        
                        try:
                            # Try to load the page
                            response = page.goto(url, wait_until='load', timeout=15000)
                            
                            if not response:
                                console.print(f"[red]No response from {url}[/red]")
                                continue
                                
                            if response.status >= 400:
                                console.print(f"[red]Error {response.status} from {url}[/red]")
                                continue
                            
                            # Wait for content
                            page.wait_for_selector('body', timeout=5000)
                            
                            # Get page height
                            height = page.evaluate('document.documentElement.scrollHeight')
                            if height < 1080:
                                height = 1080
                                
                            # Set viewport size
                            page.set_viewport_size({'width': 1920, 'height': height})
                            
                            # Take screenshot
                            output_path = os.path.join(self.temp_dir, f"{domain}_{protocol}.png")
                            page.screenshot(
                                path=output_path,
                                full_page=True,
                                type='png'
                            )
                            
                            # Verify screenshot
                            with Image.open(output_path) as img:
                                if img.size[0] >= 1920 and img.size[1] >= 1080:
                                    self.ip_data[ip]['screenshots'].add(output_path)
                                    success = True
                                    console.print(f"[green]Successfully captured screenshot of {url}[/green]")
                                    break
                                else:
                                    console.print(f"[red]Screenshot too small: {img.size}[/red]")
                                    os.remove(output_path)
                            
                        except PlaywrightTimeout:
                            console.print(f"[red]Timeout loading {url}[/red]")
                        except Exception as e:
                            console.print(f"[red]Error capturing {url}: {str(e)}[/red]")
                    
                    # Clean up
                    context.close()
                    browser.close()
                    
                except Exception as e:
                    console.print(f"[red]Browser error: {str(e)}[/red]")
                
        except Exception as e:
            console.print(f"[red]Screenshot error: {str(e)}[/red]")
        
        if not success:
            console.print(f"[red]Failed to capture screenshot for {domain}[/red]")

    def create_word_document(self, output_file: str = "report.docx"):
        doc = Document()
        
        # Create table with the specified headers
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        headers = ['IP Address', 'Domain Name(s)', 'Service(s)', 'Open Ports', 'Screenshot']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True

        # Add data rows
        for ip, data in self.ip_data.items():
            row_cells = table.add_row().cells
            
            # IP Address
            row_cells[0].text = ip
            
            # Domain Names
            row_cells[1].text = '\n'.join(sorted(data['domains']))
            
            # Services
            row_cells[2].text = '\n'.join(sorted(data['services']))
            
            # Open Ports
            row_cells[3].text = '\n'.join(str(port) for port in sorted(data['ports']))
            
            # Screenshots
            if data['screenshots']:
                paragraph = row_cells[4].paragraphs[0]
                for screenshot_path in sorted(data['screenshots']):
                    if os.path.exists(screenshot_path):
                        try:
                            # Resize image
                            img = Image.open(screenshot_path)
                            aspect_ratio = img.width / img.height
                            new_height = 200
                            new_width = int(aspect_ratio * new_height)
                            img = img.resize((new_width, new_height))
                            
                            # Save resized image
                            resized_path = screenshot_path.replace('.png', '_resized.png')
                            img.save(resized_path)
                            
                            # Add screenshot
                            paragraph.add_run().add_picture(
                                resized_path,
                                height=Inches(2)
                            )
                        except Exception as e:
                            console.print(f"[red]Error processing screenshot: {str(e)}[/red]")
            else:
                row_cells[4].text = "No web service detected"

        doc.save(output_file)
        console.print(f"\n[green]Report saved as {output_file}[/green]")

def main():
    import argparse
    
    # Create argument parser
    parser = argparse.ArgumentParser(
        description="""
Domain Analyzer - A tool for analyzing domains and IP addresses.
        
This tool performs:
- Subdomain enumeration
- Port scanning
- Service detection
- Screenshot capture
- Report generation
        """,
        formatter_class=argparse.RawDescriptionHelpFormatter
    )
    
    # Add arguments
    parser.add_argument(
        'input_file',
        help='Input file containing domains/IPs (one per line)'
    )
    
    parser.add_argument(
        '-o', '--output',
        default='report.docx',
        help='Output report filename (default: report.docx)'
    )
    
    parser.add_argument(
        '-q', '--quiet',
        action='store_true',
        help='Suppress non-essential output'
    )
    
    parser.add_argument(
        '-t', '--threads',
        type=int,
        default=min(32, (os.cpu_count() or 1) * 4),
        help='Maximum number of concurrent threads (default: CPU count * 4)'
    )
    
    parser.add_argument(
        '--timeout',
        type=int,
        default=30,
        help='Timeout in seconds for network operations (default: 30)'
    )
    
    # Parse arguments
    args = parser.parse_args()
    
    # Validate input file
    if not os.path.exists(args.input_file):
        console.print(f"[red]Error: File '{args.input_file}' not found[/red]")
        sys.exit(1)
        
    try:
        with open(args.input_file, 'r') as f:
            content = f.read().strip()
            if not content:
                console.print("[red]Error: Input file is empty[/red]")
                sys.exit(1)
                
            # Check if file contains at least one valid domain or IP
            lines = content.split('\n')
            valid_entries = False
            for line in lines:
                if DomainAnalyzer.is_valid_ip(line.strip()) or validators.domain(line.strip()):
                    valid_entries = True
                    break
                    
            if not valid_entries:
                console.print("[red]Error: No valid domains or IPs found in file[/red]")
                sys.exit(1)
                
    except Exception as e:
        console.print(f"[red]Error reading file: {str(e)}[/red]")
        sys.exit(1)
    
    # Show banner unless quiet mode is enabled
    if not args.quiet:
        console.print(Panel.fit(
            "[bold blue]Domain Analyzer[/bold blue]\n"
            "A tool for analyzing domains and IP addresses",
            border_style="bold blue"
        ))
    
    # Initialize analyzer with arguments
    analyzer = DomainAnalyzer(
        args.input_file,
        max_workers=args.threads,
        timeout=args.timeout,
        quiet=args.quiet
    )
    
    try:
        asyncio.run(analyzer.analyze_async())
    except KeyboardInterrupt:
        console.print("\n[bold red]Analysis interrupted by user[/bold red]")
        sys.exit(1)
    except Exception as e:
        console.print(f"\n[bold red]Error during analysis: {str(e)}[/bold red]")
        sys.exit(1)

if __name__ == "__main__":
    main() 